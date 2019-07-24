from docx import Document
from docx.shared import Inches
import json

def to_word(obj):
  document = Document()
  new_obj = json.loads(obj)

  document.add_heading('TLDR', 0)
  for i in new_obj:
    document.add_heading(i['heading'])
    document.add_paragraph(i['body'])
  document.save('demo.docx')

sample_data = '[ {"heading": "Heading 1", "body":"Lorem ipsum dolor sit amet, consectetur adipiscing elit. Aenean accumsan, est tempus rhoncus scelerisque, nisi eros malesuada lacus, at lobortis risus ex ac urna. Integer pulvinar congue urna tincidunt pharetra. Cras mattis ex sit amet enim rhoncus lacinia. Vivamus imperdiet porttitor sagittis."} , {"heading": "Heading 2", "body":"Lorem ipsum dolor sit amet, consectetur adipiscing elit. Aenean accumsan, est tempus rhoncus scelerisque, nisi eros malesuada lacus, at lobortis risus ex ac urna. Integer pulvinar congue urna tincidunt pharetra. Cras mattis ex sit amet enim rhoncus lacinia. Vivamus imperdiet porttitor sagittis."},{"heading": "Heading 3", "body":"Lorem ipsum dolor sit amet, consectetur adipiscing elit. Aenean accumsan, est tempus rhoncus scelerisque, nisi eros malesuada lacus, at lobortis risus ex ac urna. Integer pulvinar congue urna tincidunt pharetra. Cras mattis ex sit amet enim rhoncus lacinia. Vivamus imperdiet porttitor sagittis."}, {"heading": "Heading 4", "body":"Lorem ipsum dolor sit amet, consectetur adipiscing elit. Aenean accumsan, est tempus rhoncus scelerisque, nisi eros malesuada lacus, at lobortis risus ex ac urna. Integer pulvinar congue urna tincidunt pharetra. Cras mattis ex sit amet enim rhoncus lacinia. Vivamus imperdiet porttitor sagittis."} ] '
to_word(sample_data)